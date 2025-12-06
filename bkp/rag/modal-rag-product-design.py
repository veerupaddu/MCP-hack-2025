"""
Extended Modal RAG to include product design documents (Word, PDF, Excel)
This extends the existing modal-rag.py to support querying the product design spec
"""

import modal

app = modal.App("insurance-rag-product-design")

# Reference your specific volume
vol = modal.Volume.from_name("mcp-hack-ins-products", create_if_missing=True)

# Model configuration (same as existing)
LLM_MODEL = "microsoft/Phi-3-mini-4k-instruct"
EMBEDDING_MODEL = "BAAI/bge-small-en-v1.5"

# Build image with ALL required dependencies
image = (
    modal.Image.debian_slim(python_version="3.11")
    .pip_install(
        "vllm==0.6.3.post1",
        "langchain==0.3.7",
        "langchain-community==0.3.7",
        "langchain-text-splitters==0.3.2",
        "sentence-transformers==3.3.0",
        "chromadb==0.5.20",
        "pypdf==5.1.0",  # For PDF documents
        "python-docx==1.1.0",  # For Word documents
        "openpyxl==3.1.2",  # For Excel documents (.xlsx)
        "pandas==2.2.0",  # For Excel data processing
        "xlrd==2.0.1",  # For older Excel files (.xls)
        "cryptography==43.0.3",
        "transformers==4.46.2",
        "torch==2.4.0",
        "huggingface_hub==0.26.2",
    )
)

@app.function(image=image, volumes={"/insurance-data": vol})
def list_volume_files():
    """List all files in the volume to debug"""
    import os
    print("🔍 Listing all files in /insurance-data...")
    all_files = []
    for root, dirs, files in os.walk("/insurance-data"):
        for file in files:
            full_path = os.path.join(root, file)
            all_files.append(full_path)
            print(f"  📄 {full_path}")
    return all_files

@app.function(image=image, volumes={"/insurance-data": vol})
def load_product_design_docs():
    """Load product design documents (Word, PDF, Excel only - no markdown)"""
    import os
    import docx
    from pathlib import Path
    
    documents = []
    
    # First, list what's actually in the volume for debugging
    print("🔍 Scanning volume for product design documents (Word, PDF, Excel only)...")
    all_files = []
    for root, dirs, files in os.walk("/insurance-data"):
        for file in files:
            full_path = os.path.join(root, file)
            all_files.append(full_path)
            # Only show supported file types
            file_lower = file.lower()
            if file.endswith(('.docx', '.pdf', '.xlsx', '.xls')):
                # Relaxed filter: include all supported documents
                print(f"  📄 Found: {full_path}")
    
    # Load PDF files
    pdf_files = []
    for root, dirs, files in os.walk("/insurance-data"):
        for file in files:
            if file.endswith('.pdf'):
                full_path = os.path.join(root, file)
                file_lower = file.lower()
                # Relaxed filter: include all PDFs
                pdf_files.append(full_path)
    
    print(f"📄 Found {len(pdf_files)} PDF product design files")
    for pdf_file in pdf_files:
        try:
            from pypdf import PdfReader
            reader = PdfReader(pdf_file)
            text_content = []
            for page in reader.pages:
                text_content.append(page.extract_text())
            
            full_text = '\n'.join(text_content)
            if not full_text.strip():
                print(f"  ⚠️ No text extracted from {pdf_file}")
                continue
            
            documents.append({
                'page_content': full_text,
                'metadata': {
                    'source': pdf_file,
                    'type': 'product_design',
                    'format': 'pdf'
                }
            })
            print(f"  ✅ Loaded: {pdf_file} ({len(full_text)} characters)")
        except Exception as e:
            print(f"  ⚠️ Error loading {pdf_file}: {e}")
    
    # Load Excel files
    excel_files = []
    for root, dirs, files in os.walk("/insurance-data"):
        for file in files:
            if file.endswith(('.xlsx', '.xls')):
                full_path = os.path.join(root, file)
                file_lower = file.lower()
                if 'tokyo_auto_insurance' in file_lower or 'product_design' in file_lower:
                    excel_files.append(full_path)
    
    print(f"📊 Found {len(excel_files)} Excel product design files")
    for excel_file in excel_files:
        try:
            import pandas as pd
            # Read all sheets
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            text_content = []
            
            for sheet_name, df in excel_data.items():
                text_content.append(f"Sheet: {sheet_name}")
                # Convert DataFrame to text representation
                text_content.append(df.to_string())
                text_content.append("")  # Empty line between sheets
            
            full_text = '\n'.join(text_content)
            if not full_text.strip():
                print(f"  ⚠️ No data extracted from {excel_file}")
                continue
            
            documents.append({
                'page_content': full_text,
                'metadata': {
                    'source': excel_file,
                    'type': 'product_design',
                    'format': 'excel'
                }
            })
            print(f"  ✅ Loaded: {excel_file} ({len(full_text)} characters)")
        except Exception as e:
            print(f"  ⚠️ Error loading {excel_file}: {e}")
    
    # Load Word documents - check both root and docs subdirectory
    docx_files = []
    for root, dirs, files in os.walk("/insurance-data"):
        for file in files:
            if file.endswith('.docx'):
                full_path = os.path.join(root, file)
                # Match product design files (case insensitive)
                file_lower = file.lower()
                if 'product_design' in file_lower or 'tokyo_auto_insurance' in file_lower:
                    docx_files.append(full_path)
    
    print(f"📄 Found {len(docx_files)} Word product design files")
    for docx_file in docx_files:
        try:
            # Check if file exists
            if not os.path.exists(docx_file):
                print(f"  ⚠️ File does not exist: {docx_file}")
                continue
            
            # Check file size
            file_size = os.path.getsize(docx_file)
            print(f"  📊 File size: {file_size} bytes")
            
            # Try opening with python-docx
            # python-docx might have issues with Modal volume files, so we'll try a workaround
            try:
                doc = docx.Document(docx_file)
            except Exception as e1:
                # If direct opening fails, try copying to temp first
                import tempfile
                import shutil
                print(f"  ⚠️ Direct open failed: {e1}, trying temp copy...")
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                    shutil.copy2(docx_file, tmp.name)
                    doc = docx.Document(tmp.name)
                    tmp_path = tmp.name
                
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text_content.append(row_text)
            
            full_text = '\n'.join(text_content)
            if not full_text.strip():
                print(f"  ⚠️ No text extracted from {docx_file}")
                continue
                
            documents.append({
                'page_content': full_text,
                'metadata': {
                    'source': docx_file,
                    'type': 'product_design',
                    'format': 'word'
                }
            })
            print(f"  ✅ Loaded: {docx_file} ({len(full_text)} characters)")
            
            # Clean up temp file if we created one
            if 'tmp_path' in locals():
                try:
                    os.unlink(tmp_path)
                except:
                    pass
                    
        except Exception as e:
            print(f"  ⚠️ Error loading {docx_file}: {e}")
            import traceback
            print(f"  Traceback: {traceback.format_exc()}")
    
    print(f"✅ Loaded {len(documents)} product design documents")
    return documents

@app.function(
    image=image,
    volumes={"/insurance-data": vol},
    gpu="T4",  # Enable GPU for faster embedding generation
    timeout=1800  # Increased timeout to 30 minutes
)
def create_product_design_vector_db():
    """Create vector database from product design documents"""
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    
    print("🔍 Loading product design documents...")
    documents = load_product_design_docs.remote()
    
    if len(documents) == 0:
        return {
            "status": "error",
            "message": "No product design documents found",
            "total_documents": 0,
            "total_chunks": 0
        }
    
    print("✂️ Splitting documents into chunks...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    
    # Convert to LangChain document format
    from langchain.schema import Document
    langchain_docs = []
    for doc in documents:
        langchain_docs.append(Document(
            page_content=doc['page_content'],
            metadata=doc['metadata']
        ))
    
    chunks = text_splitter.split_documents(langchain_docs)
    print(f"📦 Created {len(chunks)} chunks")
    
    print("🧠 Creating embeddings...")
    # Try CUDA first, fall back to CPU if not available
    import torch
    device = 'cuda' if torch.cuda.is_available() else 'cpu'
    print(f"   Using device: {device}")
    
    embeddings = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={'device': device},
        encode_kwargs={'normalize_embeddings': True}
    )
    
    print("💾 Building vector database...")
    
    # Connect to remote Chroma service
    chroma_service = modal.Cls.from_name("chroma-server-v2", "ChromaDB")()
    
    # Prepare data for upsert
    ids = [f"product_design_{i}" for i in range(len(chunks))]
    documents_text = [chunk.page_content for chunk in chunks]
    metadatas = [chunk.metadata for chunk in chunks]
    
    # Generate embeddings locally
    print("   Generating embeddings locally...")
    embeddings_list = embeddings.embed_documents(documents_text)
    
    # Upsert to remote Chroma (use separate collection for product design)
    print("   Upserting to remote Chroma DB...")
    batch_size = 100
    for i in range(0, len(ids), batch_size):
        batch_ids = ids[i:i+batch_size]
        batch_docs = documents_text[i:i+batch_size]
        batch_metas = metadatas[i:i+batch_size]
        batch_embs = embeddings_list[i:i+batch_size]
        
        chroma_service.upsert.remote(
            collection_name="product_design",  # Separate collection
            ids=batch_ids,
            documents=batch_docs,
            embeddings=batch_embs,
            metadatas=batch_metas
        )
        print(f"   Upserted batch {i//batch_size + 1}/{(len(ids)-1)//batch_size + 1}")
    
    print("✅ Product design vector database created!")
    
    return {
        "status": "success",
        "total_documents": len(documents),
        "total_chunks": len(chunks)
    }

@app.cls(
    image=image,
    volumes={"/insurance-data": vol},
    gpu="A10G",
    timeout=600,
    max_containers=1,
    min_containers=0
)
class ProductDesignRAG:
    """RAG model specifically for product design document queries"""
    
    @modal.enter()
    def enter(self):
        from langchain_community.embeddings import HuggingFaceEmbeddings
        from vllm import LLM, SamplingParams
        from langchain.schema import Document
        
        print("🚀 Initializing Product Design RAG...")
        
        # Initialize embeddings
        self.embeddings = HuggingFaceEmbeddings(
            model_name=EMBEDDING_MODEL,
            model_kwargs={'device': 'cuda'},
            encode_kwargs={'normalize_embeddings': True}
        )
        
        # Connect to Chroma
        self.chroma_service = modal.Cls.from_name("chroma-server-v2", "ChromaDB")()
        
        # Custom retriever for remote Chroma
        class RemoteChromaRetriever:
            def __init__(self, chroma_service, embeddings, k=3):
                self.chroma_service = chroma_service
                self.embeddings = embeddings
                self.k = k
            
            def get_relevant_documents(self, query: str):
                query_embedding = self.embeddings.embed_query(query)
                results = self.chroma_service.query.remote(
                    collection_name="product_design",
                    query_embeddings=[query_embedding],
                    n_results=self.k
                )
                
                docs = []
                if results and 'documents' in results and len(results['documents']) > 0:
                    for i, doc_text in enumerate(results['documents'][0]):
                        metadata = results.get('metadatas', [[{}]])[0][i] if 'metadatas' in results else {}
                        docs.append(Document(page_content=doc_text, metadata=metadata))
                
                return docs
        
        self.RemoteChromaRetriever = RemoteChromaRetriever
        
        # Load LLM
        print("   Loading LLM...")
        self.llm_engine = LLM(
            model=LLM_MODEL,
            dtype="float16",
            gpu_memory_utilization=0.85,
            max_model_len=4096,
            trust_remote_code=True,
            enforce_eager=True
        )
        
        self.sampling_params = SamplingParams(
            temperature=0.7,
            max_tokens=1536,  # Increased for comprehensive, detailed answers
            top_p=0.9,
            stop=["\n\n\n", "Question:", "Context:", "<|end|>"]  # Removed single \n\n to allow longer answers
        )
        
        print("✅ Product Design RAG ready!")
    
    @modal.method()
    def query(self, question: str, top_k: int = 5):  # Increased from 3 to 5 for more context
        """Query the product design document"""
        import time
        start_time = time.time()
        
        print(f"❓ Query: {question}")
        
        # Retrieve relevant documents
        retrieval_start = time.time()
        retriever = self.RemoteChromaRetriever(
            chroma_service=self.chroma_service,
            embeddings=self.embeddings,
            k=top_k
        )
        docs = retriever.get_relevant_documents(question)
        retrieval_time = time.time() - retrieval_start
        
        if not docs:
            return {
                "question": question,
                "answer": "No relevant information found in the product design document.",
                "retrieval_time": retrieval_time,
                "generation_time": 0,
                "sources": []
            }
        
        # Build context
        context = "\n\n".join([doc.page_content for doc in docs])
        
        # Create prompt with instructions for comprehensive answers
        prompt = f"""<|system|>
You are a helpful AI assistant that answers questions about the provided insurance product documents. 
Provide comprehensive, detailed answers with specific information from the document. 
Structure your answer clearly with:
- A brief summary if relevant
- Detailed explanations with specific numbers, percentages, and data points
- Step-by-step guidance when appropriate
- Clear formatting (use bullet points or numbered lists when helpful)
Be thorough and cite specific details from the context. If information is not available, say so clearly.<|end|>
<|user|>
Context from Product Design Document:
{context}

Question:
{question}<|end|>
<|assistant|>"""
        
        # Generate answer
        outputs = self.llm_engine.generate(prompts=[prompt], sampling_params=self.sampling_params)
        answer = outputs[0].outputs[0].text.strip()
        generation_time = time.time() - start_time - retrieval_time
        
        # Prepare sources with better content extraction
        sources = []
        for doc in docs:
            # Clean up source content - remove markdown table syntax
            content = doc.page_content
            # Remove markdown table separators
            import re
            content = re.sub(r'\|[\s\-:]+\|', '', content)
            content = re.sub(r'^\|.*\|$', '', content, flags=re.MULTILINE)
            content = re.sub(r'\s+\|\s+', ' ', content)
            content = content.strip()
            
            sources.append({
                "content": content[:500],  # Increased from 300 to 500
                "metadata": doc.metadata
            })
        
        return {
            "question": question,
            "answer": answer,
            "retrieval_time": retrieval_time,
            "generation_time": generation_time,
            "sources": sources
        }

@app.local_entrypoint()
def list_files():
    """List files in volume for debugging"""
    print("🔍 Listing files in volume...")
    files = list_volume_files.remote()
    print(f"\n✅ Found {len(files)} files total")

@app.local_entrypoint()
def index_product_design():
    """Index product design documents"""
    print("🚀 Indexing product design documents...")
    
    # First, list files to debug
    print("\n📋 Checking volume contents...")
    try:
        files = list_volume_files.remote()
        print(f"Found {len(files)} files in volume")
    except Exception as e:
        print(f"Could not list files: {e}")
    
    result = create_product_design_vector_db.remote()
    print(f"\n{'='*60}")
    print(f"Status: {result['status']}")
    if result['status'] == 'success':
        print(f"Documents processed: {result['total_documents']}")
        print(f"Text chunks created: {result['total_chunks']}")
        print("✅ Product design vector database is ready!")
    else:
        print(f"❌ Error: {result['message']}")
        print("\n💡 Tip: Make sure files are uploaded to the volume:")
        print("   modal volume put mcp-hack-ins-products \\")
        print("     docs/tokyo_auto_insurance_product_design.docx")
        print("   # Or PDF/Excel files:")
        print("   # modal volume put mcp-hack-ins-products docs/file.pdf")
        print("   # modal volume put mcp-hack-ins-products docs/file.xlsx")
    print(f"{'='*60}")

@app.local_entrypoint()
def query_product_design(question: str = "What are the three product tiers and their premium ranges?"):
    """Query the product design document"""
    print(f"🤔 Question: {question}\n")
    
    model = ProductDesignRAG()
    result = model.query.remote(question)
    
    print(f"{'='*60}")
    print(f"📝 Answer:")
    print(f"{result['answer']}\n")
    print(f"{'='*60}")
    print(f"⏱️  Retrieval: {result['retrieval_time']:.2f}s")
    print(f"⏱️  Generation: {result['generation_time']:.2f}s")
    
    if result['sources']:
        print(f"\n📚 Sources ({len(result['sources'])}):")
        for i, source in enumerate(result['sources'], 1):
            print(f"\n{i}. {source['metadata'].get('source', 'Unknown')}")
            print(f"   {source['content'][:200]}...")

# Define data model for API
from pydantic import BaseModel

class RAGQuery(BaseModel):
    question: str
    top_k: int = 5

@app.function(image=image)
@modal.web_endpoint(method="POST")
def api_query(item: RAGQuery):
    """Expose RAG query as a web endpoint"""
    model = ProductDesignRAG()
    return model.query.remote(item.question, item.top_k)

