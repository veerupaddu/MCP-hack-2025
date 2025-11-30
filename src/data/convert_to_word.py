#!/usr/bin/env python3
"""
Convert markdown to Word document using python-docx
"""

import sys
import subprocess
import re

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_word(md_file, docx_file):
    """Convert markdown file to Word document"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines (but add spacing)
        if not line:
            i += 1
            continue
        
        # Title (first line with #)
        if line.startswith('# ') and i < 5:
            title = line[2:].strip()
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Headings
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue
        
        # Tables
        if '|' in line and line.count('|') >= 2:
            # Collect table rows
            table_rows = []
            j = i
            while j < len(lines) and '|' in lines[j] and lines[j].count('|') >= 2:
                if not lines[j].strip().startswith('|---'):
                    table_rows.append(lines[j])
                j += 1
            
            if table_rows:
                # Parse table
                headers = [cell.strip() for cell in table_rows[0].split('|')[1:-1]]
                data_rows = []
                for row in table_rows[1:]:
                    cells = [cell.strip() for cell in row.split('|')[1:-1]]
                    if cells:
                        data_rows.append(cells)
                
                # Create table
                if headers:
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    header_cells = table.rows[0].cells
                    for idx, header in enumerate(headers):
                        header_cells[idx].text = header
                        header_cells[idx].paragraphs[0].runs[0].font.bold = True
                    
                    # Data rows
                    for row_data in data_rows:
                        if len(row_data) == len(headers):
                            row_cells = table.add_row().cells
                            for idx, cell_data in enumerate(row_data):
                                row_cells[idx].text = cell_data
                
                i = j
                continue
        
        # Lists
        if line.startswith('- ') or line.startswith('* '):
            items = []
            j = i
            while j < len(lines) and (lines[j].strip().startswith('- ') or 
                                      lines[j].strip().startswith('* ') or
                                      lines[j].strip().startswith('  - ') or
                                      lines[j].strip().startswith('  * ')):
                items.append(lines[j].strip()[2:].strip())
                j += 1
            
            for item in items:
                # Remove markdown formatting
                item = re.sub(r'\*\*(.*?)\*\*', r'\1', item)
                item = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', item)
                p = doc.add_paragraph(item, style='List Bullet')
            
            i = j
            continue
        
        # Regular paragraph
        # Remove markdown formatting
        text = line
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        
        if text.strip():
            p = doc.add_paragraph(text)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✓ Word document created: {docx_file}")

if __name__ == "__main__":
    md_file = "docs/tokyo_auto_insurance_product_design_filled.md"
    docx_file = "docs/tokyo_auto_insurance_product_design.docx"
    
    try:
        markdown_to_word(md_file, docx_file)
    except Exception as e:
        print(f"Error: {e}")
        print("\nTrying alternative: Creating a simpler Word document...")
        # Fallback: create a basic Word doc
        from docx import Document
        doc = Document()
        doc.add_heading('Auto Insurance Product Design Document for Tokyo, Japan', 0)
        
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Simple conversion - just paragraphs
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                doc.add_heading(line.lstrip('# '), level=min(level, 3))
            else:
                # Remove markdown
                line = line.replace('**', '')
                line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', line)
                if line:
                    doc.add_paragraph(line)
        
        doc.save(docx_file)
        print(f"✓ Word document created (simple format): {docx_file}")

