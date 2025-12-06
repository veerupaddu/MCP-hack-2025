"""
RAG Indexing for New Product Design Documents
Indexes product design documents from docs/ folder into 'product_design' collection.

Folders indexed:
- /insurance-data/docs/
- /insurance-data/docs/product-design/

Supports: .docx, .xlsx, .xls files
"""

import modal

app = modal.App("insurance-rag-product-design-v2")

# Reference the volume
vol = modal.Volume.from_name("mcp-hack-ins-products", create_if_missing=True)

# Model configuration
EMBEDDING_MODEL = "BAAI/bge-small-en-v1.5"

# Source folders to index
SOURCE_FOLDERS = [
    "/insurance-data/docs",
    "/insurance-data/docs/product-design"
]

# Build image
image = (
    modal.Image.debian_slim(python_version="3.11")
    .pip_install(
        "torch>=2.0.0",
        "transformers>=4.30.0",
        "sentence-transformers>=2.2.0",
        "huggingface_hub>=0.15.0",
        "langchain>=0.1.0",
        "langchain-community>=0.0.13",
        "langchain-text-splitters>=0.2.0",
        "langchain-core>=0.1.0",
        "langchain-huggingface>=0.0.3",  # New HuggingFace integration
        "chromadb>=0.4.0",
        "python-docx>=1.1.0",
        "openpyxl>=3.1.0",
        "pandas>=2.0.0",
        "xlrd>=2.0.0",
    )
)


@app.function(image=image, volumes={"/insurance-data": vol})
def list_product_design_files():
    """List all product design files"""
    import os
    
    print("🔍 Scanning product design folders...")
    all_files = []
    
    for folder in SOURCE_FOLDERS:
        if os.path.exists(folder):
            print(f"\n📁 {folder}:")
            for file in os.listdir(folder):
                if file.endswith(('.docx', '.xlsx', '.xls')):
                    full_path = os.path.join(folder, file)
                    all_files.append(full_path)
                    print(f"  📄 {file}")
        else:
            print(f"  ⚠️ Folder not found: {folder}")
    
    print(f"\n✅ Found {len(all_files)} product design files total")
    return all_files


@app.function(image=image, volumes={"/insurance-data": vol}, timeout=600)
def load_product_design_docs():
    """Load all product design documents (DOCX, XLSX)"""
    import os
    import docx
    import pandas as pd
    
    documents = []
    
    print("📚 Loading product design documents...")
    
    for folder in SOURCE_FOLDERS:
        if not os.path.exists(folder):
            print(f"  ⚠️ Skipping non-existent folder: {folder}")
            continue
            
        print(f"\n📁 Processing {folder}/")
        
        for file in os.listdir(folder):
            full_path = os.path.join(folder, file)
            
            # Skip directories
            if os.path.isdir(full_path):
                continue
            
            try:
                if file.endswith('.docx'):
                    # Load Word document
                    doc = docx.Document(full_path)
                    text_content = []
                    
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_content.append(para.text)
                    
                    # Extract tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' | '.join([cell.text for cell in row.cells])
                            if row_text.strip():
                                text_content.append(row_text)
                    
                    full_text = '\n'.join(text_content)
                    doc_type = 'word'
                    
                elif file.endswith(('.xlsx', '.xls')):
                    # Load Excel document
                    excel_data = pd.read_excel(full_path, sheet_name=None)
                    text_content = []
                    
                    for sheet_name, df in excel_data.items():
                        text_content.append(f"Sheet: {sheet_name}")
                        text_content.append(df.to_string())
                        text_content.append("")
                    
                    full_text = '\n'.join(text_content)
                    doc_type = 'excel'
                    
                else:
                    continue
                
                if not full_text.strip():
                    print(f"  ⚠️ No text extracted: {file}")
                    continue
                
                documents.append({
                    'page_content': full_text,
                    'metadata': {
                        'source': full_path,
                        'filename': file,
                        'type': 'product_design',
                        'format': doc_type
                    }
                })
                print(f"  ✅ {file} ({len(full_text):,} chars, {doc_type})")
                
            except Exception as e:
                print(f"  ❌ Error loading {file}: {e}")
    
    print(f"\n✅ Loaded {len(documents)} documents total")
    return documents


@app.function(
    image=image,
    volumes={"/insurance-data": vol},
    gpu="T4",
    timeout=1800
)
def index_product_design():
    """Create vector index for product design documents"""
    from langchain_huggingface import HuggingFaceEmbeddings
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document
    import chromadb
    
    print("🚀 Indexing product design documents...")
    
    # Load documents
    documents = load_product_design_docs.remote()
    
    if not documents:
        return {
            "status": "error",
            "message": "No documents found to index",
            "total_documents": 0,
            "total_chunks": 0
        }
    
    # Split into chunks
    print("\n✂️ Splitting documents into chunks...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    
    langchain_docs = [
        Document(page_content=doc['page_content'], metadata=doc['metadata'])
        for doc in documents
    ]
    
    chunks = text_splitter.split_documents(langchain_docs)
    print(f"📦 Created {len(chunks)} chunks")
    
    # Create embeddings
    print("\n🧠 Creating embeddings...")
    import torch
    device = 'cuda' if torch.cuda.is_available() else 'cpu'
    print(f"   Using device: {device}")
    
    embeddings = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={'device': device},
        encode_kwargs={'normalize_embeddings': True}
    )
    
    # Store in ChromaDB
    print("\n💾 Storing in ChromaDB (product_design collection)...")
    chroma_client = chromadb.PersistentClient(path="/insurance-data/chroma_db")
    
    # Delete existing collection if exists
    try:
        chroma_client.delete_collection("product_design")
        print("   Deleted old collection")
    except:
        pass
    
    collection = chroma_client.create_collection(
        name="product_design",
        metadata={"description": "TokyoDrive product design specifications"}
    )
    
    # Generate embeddings and upsert
    batch_size = 50
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i+batch_size]
        
        ids = [f"design_{i+j}" for j in range(len(batch))]
        documents_text = [chunk.page_content for chunk in batch]
        metadatas = [chunk.metadata for chunk in batch]
        embeddings_list = embeddings.embed_documents(documents_text)
        
        collection.add(
            ids=ids,
            documents=documents_text,
            embeddings=embeddings_list,
            metadatas=metadatas
        )
        print(f"   Batch {i//batch_size + 1}/{(len(chunks)-1)//batch_size + 1} complete")
    
    # Commit to volume
    vol.commit()
    
    print("\n✅ Indexing complete!")
    
    return {
        "status": "success",
        "collection": "product_design",
        "total_documents": len(documents),
        "total_chunks": len(chunks),
        "files": [doc['metadata']['filename'] for doc in documents]
    }


@app.local_entrypoint()
def main():
    """Run the indexing process"""
    print("=" * 60)
    print("Product Design Documents Indexer")
    print("=" * 60)
    
    # First list files
    print("\n📋 Checking available files...")
    files = list_product_design_files.remote()
    
    if not files:
        print("\n❌ No product design files found!")
        print("   Make sure files are uploaded to the volume:")
        for folder in SOURCE_FOLDERS:
            print(f"   - {folder}")
        return
    
    # Index documents
    print("\n🚀 Starting indexing...")
    result = index_product_design.remote()
    
    print("\n" + "=" * 60)
    print("Result:")
    print(f"  Status: {result['status']}")
    if result['status'] == 'success':
        print(f"  Documents indexed: {result['total_documents']}")
        print(f"  Chunks created: {result['total_chunks']}")
        print(f"  Files: {', '.join(result['files'])}")
    else:
        print(f"  Error: {result.get('message', 'Unknown error')}")
    print("=" * 60)
